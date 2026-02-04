from ui import run_app


if __name__ == "__main__":
    run_app()
    raise SystemExit()

"""
import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF
from pdf2docx.converter import Converter
try:
    from docx2pdf import convert as docx_to_pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    DOCX2PDF_AVAILABLE = False
import threading
from io import BytesIO
import traceback
import shutil
from copy import deepcopy

try:
    from tkinterdnd2 import DND_FILES, TkinterDnD
    DRAG_DROP_AVAILABLE = True
except ImportError:
    DRAG_DROP_AVAILABLE = False


class ConsolidadorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Consolidador de Documentos - Arrastrar y Soltar")
        self.root.geometry("750x700")
        self.root.resizable(True, True)
        
        self.archivos_seleccionados = []
        self.carpeta_entrada = None
        self.agregar_pagina_blanca = tk.BooleanVar(value=False)
        self.normalizar_a_pdf = tk.BooleanVar(value=False)
        self.formato_salida = tk.StringVar(value="DOCX")
        
        # Frame principal
        main_frame = ttk.Frame(root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Título
        titulo = ttk.Label(main_frame, text="Consolidador de Documentos (Word y PDF)",
                  font=("Arial", 14, "bold"))
        titulo.grid(row=0, column=0, columnspan=3, pady=10)
        
        # Instrucciones
        instrucciones = ttk.Label(main_frame, 
                                 text="✨ Arrastra archivos PDF/DOCX aquí o usa los botones ✨",
                                 font=("Arial", 10, "italic"),
                                 foreground="blue")
        instrucciones.grid(row=1, column=0, columnspan=3, pady=5)
        
        # Frame de botones
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=2, column=0, columnspan=3, pady=10, sticky=(tk.W, tk.E))
        
        # Botón agregar archivos
        self.btn_agregar = ttk.Button(btn_frame, text="➕ Agregar Archivos",
                                      command=self.agregar_archivos)
        self.btn_agregar.pack(side=tk.LEFT, padx=5)
        
        # Botón seleccionar carpeta
        self.btn_carpeta = ttk.Button(btn_frame, text="📁 Cargar Carpeta",
                                      command=self.seleccionar_carpeta)
        self.btn_carpeta.pack(side=tk.LEFT, padx=5)
        
        # Botón limpiar lista
        self.btn_limpiar = ttk.Button(btn_frame, text="🗑️ Limpiar Lista",
                                      command=self.limpiar_lista)
        self.btn_limpiar.pack(side=tk.LEFT, padx=5)
        
        # Lista de archivos
        ttk.Label(main_frame, text="Archivos en orden de consolidación:",
                 font=("Arial", 10, "bold")).grid(row=3, column=0, 
                                                 columnspan=3, pady=(10, 5))
        
        # Frame para listbox y botones de orden
        listbox_frame = ttk.Frame(main_frame)
        listbox_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)
        
        # Listbox con scroll
        scroll_frame = ttk.Frame(listbox_frame)
        scroll_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        self.listbox = tk.Listbox(scroll_frame, height=12, selectmode=tk.SINGLE,
                                  bg="#f5f5f5", font=("Consolas", 9),
                                  activestyle='dotbox')
        self.listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Scrollbar para listbox
        scrollbar = ttk.Scrollbar(scroll_frame, orient=tk.VERTICAL, command=self.listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.listbox.config(yscrollcommand=scrollbar.set)
        
        # Configurar drag and drop si está disponible
        if DRAG_DROP_AVAILABLE:
            self.listbox.drop_target_register(DND_FILES)
            self.listbox.dnd_bind('<<Drop>>', self.drop_archivos)
        
        # Frame de botones de orden
        orden_frame = ttk.Frame(listbox_frame)
        orden_frame.pack(side=tk.RIGHT, fill=tk.Y, padx=10)
        
        ttk.Label(orden_frame, text="Ordenar:", font=("Arial", 9, "bold")).pack(pady=(0, 5))
        ttk.Button(orden_frame, text="⬆️ Subir", command=self.subir_archivo, width=12).pack(pady=3)
        ttk.Button(orden_frame, text="⬇️ Bajar", command=self.bajar_archivo, width=12).pack(pady=3)
        ttk.Button(orden_frame, text="❌ Quitar", command=self.quitar_archivo, width=12).pack(pady=20)
        
        # Nombre archivo salida
        salida_frame = ttk.Frame(main_frame)
        salida_frame.grid(row=5, column=0, columnspan=3, pady=(15, 5), sticky=(tk.W, tk.E))
        
        ttk.Label(salida_frame, text="Nombre base:", 
                 font=("Arial", 10)).pack(side=tk.LEFT, padx=5)
        
        self.entry_nombre = ttk.Entry(salida_frame, width=35, font=("Arial", 10))
        self.entry_nombre.insert(0, "consolidado")
        self.entry_nombre.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        # Checkbox para páginas en blanco
        checkbox_frame = ttk.Frame(main_frame)
        checkbox_frame.grid(row=6, column=0, columnspan=3, pady=(5, 5), sticky=(tk.W))
        
        self.check_pagina = ttk.Checkbutton(
            checkbox_frame,
            text="📄 Agregar página en blanco entre cada archivo",
            variable=self.agregar_pagina_blanca,
            onvalue=True,
            offvalue=False
        )
        self.check_pagina.pack(side=tk.LEFT, padx=5)
        
        # Checkbox para normalizar a PDF primero
        checkbox_pdf_frame = ttk.Frame(main_frame)
        checkbox_pdf_frame.grid(row=7, column=0, columnspan=3, pady=(5, 5), sticky=(tk.W))
        
        self.check_normalizar = ttk.Checkbutton(
            checkbox_pdf_frame,
            text="🔄 Normalizar todo a PDF antes de consolidar (DOCX → PDF → DOCX final)",
            variable=self.normalizar_a_pdf,
            onvalue=True,
            offvalue=False,
            state='normal' if DOCX2PDF_AVAILABLE else 'disabled'
        )
        self.check_normalizar.pack(side=tk.LEFT, padx=5)
        
        if not DOCX2PDF_AVAILABLE:
            ttk.Label(checkbox_pdf_frame, text="⚠️ Instala docx2pdf", foreground="orange", font=("Arial", 8)).pack(side=tk.LEFT, padx=5)

        # Selector de formato de salida
        formato_frame = ttk.Frame(main_frame)
        formato_frame.grid(row=8, column=0, columnspan=3, pady=(5, 5), sticky=(tk.W))

        ttk.Label(formato_frame, text="Formato de salida:", font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)
        self.combo_formato = ttk.Combobox(
            formato_frame,
            textvariable=self.formato_salida,
            values=["DOCX", "PDF", "DOCX+PDF"],
            state="readonly",
            width=12
        )
        self.combo_formato.pack(side=tk.LEFT, padx=5)
        self.combo_formato.current(0)

        if not DOCX2PDF_AVAILABLE:
            ttk.Label(
                formato_frame,
                text="⚠️ PDF requiere docx2pdf",
                foreground="orange",
                font=("Arial", 8)
            ).pack(side=tk.LEFT, padx=5)
        
        # Ruta de guardado
        ruta_frame = ttk.Frame(main_frame)
        ruta_frame.grid(row=9, column=0, columnspan=3, pady=(5, 5), sticky=(tk.W, tk.E))
        
        ttk.Label(ruta_frame, text="Guardar en:", 
                 font=("Arial", 10, "bold")).pack(side=tk.LEFT, padx=5)
        
        self.entry_ruta = ttk.Entry(ruta_frame, width=40, font=("Arial", 9), state='readonly')
        self.entry_ruta.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        
        ttk.Button(ruta_frame, text="📂 Cambiar", 
                  command=self.seleccionar_ruta_salida, width=15).pack(side=tk.LEFT, padx=5)
        
        # Botón procesar
        self.btn_procesar = ttk.Button(main_frame, text="🚀 CONSOLIDAR DOCUMENTOS",
                           command=self.procesar_documentos)
        self.btn_procesar.grid(row=10, column=0, columnspan=3, pady=8, sticky=(tk.W, tk.E))

        # Botón consolidar a PDF
        self.btn_solo_pdf = ttk.Button(main_frame, text="📄 CONSOLIDAR A PDF",
                           command=self.procesar_solo_pdf)
        self.btn_solo_pdf.grid(row=11, column=0, columnspan=3, pady=8, sticky=(tk.W, tk.E))
        
        # Barra de progreso
        self.progress = ttk.Progressbar(main_frame, length=400, mode='indeterminate')
        self.progress.grid(row=12, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=5)
        
        # Label estado
        self.label_estado = ttk.Label(main_frame, text="Estado: Esperando archivos...", 
                          foreground="blue", font=("Arial", 9))
        self.label_estado.grid(row=13, column=0, columnspan=3, sticky=(tk.W, tk.E))
        
        # Configurar grid weights
        root.columnconfigure(0, weight=1)
        root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(0, weight=1)
        main_frame.rowconfigure(4, weight=1)
    
    def agregar_archivos(self):
        """Abre diálogo para seleccionar múltiples archivos"""
        archivos = filedialog.askopenfilenames(
            title="Seleccionar archivos PDF o DOCX",
            filetypes=[
                ("Documentos", "*.pdf *.docx *.doc"),
                ("PDF", "*.pdf"),
                ("Word", "*.docx *.doc"),
                ("Todos", "*.*")
            ]
        )
        
        if archivos:
            for archivo in archivos:
                if archivo not in self.archivos_seleccionados:
                    self.archivos_seleccionados.append(archivo)
            
            self.actualizar_listbox()
            self.label_estado.config(
                text=f"Estado: {len(self.archivos_seleccionados)} archivo(s) cargado(s)",
                foreground="green"
            )
    
    def drop_archivos(self, event):
        """Maneja archivos arrastrados"""
        archivos = self.root.tk.splitlist(event.data)
        
        for archivo in archivos:
            # Limpiar el path (quitar llaves si existen)
            archivo = archivo.strip('{}')
            
            if os.path.isfile(archivo):
                ext = archivo.lower()
                if ext.endswith(('.pdf', '.docx', '.doc')):
                    if archivo not in self.archivos_seleccionados:
                        self.archivos_seleccionados.append(archivo)
        
        self.actualizar_listbox()
        
        if self.archivos_seleccionados:
            self.label_estado.config(
                text=f"Estado: {len(self.archivos_seleccionados)} archivo(s) cargado(s)",
                foreground="green"
            )
    
    def limpiar_lista(self):
        """Limpia la lista de archivos"""
        if self.archivos_seleccionados:
            if messagebox.askyesno("Confirmar", "¿Limpiar toda la lista de archivos?"):
                self.archivos_seleccionados = []
                self.listbox.delete(0, tk.END)
                self.label_estado.config(text="Estado: Lista limpiada", foreground="blue")
        else:
            messagebox.showinfo("Info", "La lista ya está vacía")
    
    def subir_archivo(self):
        """Sube el archivo seleccionado en la lista"""
        seleccion = self.listbox.curselection()
        if not seleccion:
            messagebox.showinfo("Info", "Selecciona un archivo de la lista primero")
            return
        
        idx = seleccion[0]
        if idx > 0:
            # Intercambiar en lista de archivos
            self.archivos_seleccionados[idx], self.archivos_seleccionados[idx-1] = \
                self.archivos_seleccionados[idx-1], self.archivos_seleccionados[idx]
            
            # Actualizar listbox
            self.actualizar_listbox()
            self.listbox.selection_set(idx-1)
            self.listbox.see(idx-1)
    
    def bajar_archivo(self):
        """Baja el archivo seleccionado en la lista"""
        seleccion = self.listbox.curselection()
        if not seleccion:
            messagebox.showinfo("Info", "Selecciona un archivo de la lista primero")
            return
        
        idx = seleccion[0]
        if idx < len(self.archivos_seleccionados) - 1:
            # Intercambiar en lista de archivos
            self.archivos_seleccionados[idx], self.archivos_seleccionados[idx+1] = \
                self.archivos_seleccionados[idx+1], self.archivos_seleccionados[idx]
            
            # Actualizar listbox
            self.actualizar_listbox()
            self.listbox.selection_set(idx+1)
            self.listbox.see(idx+1)
    
    def quitar_archivo(self):
        """Quita el archivo seleccionado de la lista"""
        seleccion = self.listbox.curselection()
        if not seleccion:
            messagebox.showinfo("Info", "Selecciona un archivo de la lista primero")
            return
        
        idx = seleccion[0]
        nombre = os.path.basename(self.archivos_seleccionados[idx])
        
        if messagebox.askyesno("Confirmar", f"¿Quitar '{nombre}' de la lista?"):
            del self.archivos_seleccionados[idx]
            self.actualizar_listbox()
            
            self.label_estado.config(
                text=f"Estado: {len(self.archivos_seleccionados)} archivo(s) en lista",
                foreground="green"
            )
    
    def actualizar_listbox(self):
        """Actualiza la visualización del listbox"""
        self.listbox.delete(0, tk.END)
        for idx, archivo in enumerate(self.archivos_seleccionados, 1):
            nombre = os.path.basename(archivo)
            ext = os.path.splitext(nombre)[1].upper()
            self.listbox.insert(tk.END, f"{idx}. [{ext}] {nombre}")
        
        # Actualizar ruta de guardado si no está definida
        if not self.carpeta_entrada and self.archivos_seleccionados:
            self.carpeta_entrada = os.path.dirname(self.archivos_seleccionados[0])
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, self.carpeta_entrada)
            self.entry_ruta.config(state='readonly')
    
    def seleccionar_ruta_salida(self):
        """Selecciona la carpeta donde se guardará el archivo consolidado"""
        carpeta = filedialog.askdirectory(
            title="Seleccionar carpeta para guardar el consolidado",
            initialdir=self.carpeta_entrada if self.carpeta_entrada else os.path.expanduser("~")
        )
        
        if carpeta:
            self.carpeta_entrada = carpeta
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, carpeta)
            self.entry_ruta.config(state='readonly')
            self.label_estado.config(
                text=f"Estado: Ruta de guardado actualizada",
                foreground="green"
            )
    
    def seleccionar_carpeta(self):
        """Selecciona carpeta de entrada y carga todos los archivos ordenados"""
        carpeta = filedialog.askdirectory(title="Seleccionar carpeta con documentos")
        
        if carpeta:
            self.carpeta_entrada = carpeta
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, carpeta)
            self.entry_ruta.config(state='readonly')
            self.listar_archivos(carpeta)
    
    def listar_archivos(self, carpeta):
        """Lista archivos PDF y DOCX en la carpeta y los agrega a la lista"""
        # Obtener archivos
        archivos = []
        for archivo in os.listdir(carpeta):
            if archivo.lower().endswith(('.pdf', '.docx', '.doc')):
                ruta_completa = os.path.join(carpeta, archivo)
                archivos.append((archivo, ruta_completa))
        
        # Ordenar por número al inicio del nombre
        def extraer_numero(nombre_archivo):
            match = re.match(r'(\d+)', nombre_archivo[0])
            return int(match.group(1)) if match else float('inf')
        
        archivos.sort(key=extraer_numero)
        
        # Agregar a la lista existente
        agregados = 0
        for archivo, ruta in archivos:
            if ruta not in self.archivos_seleccionados:
                self.archivos_seleccionados.append(ruta)
                agregados += 1
        
        self.actualizar_listbox()
        
        if agregados > 0:
            self.label_estado.config(
                text=f"Estado: Se agregaron {agregados} archivo(s). Total: {len(self.archivos_seleccionados)}",
                foreground="green"
            )
        else:
            self.label_estado.config(
                text="Estado: No se encontraron archivos nuevos en la carpeta",
                foreground="orange"
            )
    
    def procesar_documentos(self):
        """Procesa y consolida los documentos en el orden de la lista"""
        if not self.archivos_seleccionados:
            messagebox.showwarning("Advertencia", 
                                  "No hay archivos para procesar.\n\n"
                                  "Arrastra archivos a la lista o usa los botones para agregar.")
            return
        
        nombre_base = self.entry_nombre.get().strip()
        if not nombre_base:
            nombre_base = "consolidado"

        formato = self.formato_salida.get()
        if formato == "PDF":
            self._procesar_como_pdf(nombre_base)
            return
        if formato == "DOCX+PDF" and not DOCX2PDF_AVAILABLE:
            messagebox.showwarning(
                "PDF no disponible",
                "docx2pdf no está instalado. Selecciona DOCX o instala docx2pdf."
            )
            return

        nombre_salida = f"{nombre_base}.docx"
        
        # Determinar carpeta de salida (usar la del primer archivo si no está definida)
        if not self.carpeta_entrada and self.archivos_seleccionados:
            self.carpeta_entrada = os.path.dirname(self.archivos_seleccionados[0])
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, self.carpeta_entrada)
            self.entry_ruta.config(state='readonly')
        
        # Verificar que hay una ruta de salida válida
        if not self.carpeta_entrada or not os.path.exists(self.carpeta_entrada):
            messagebox.showwarning("Advertencia", 
                                  "Por favor selecciona una carpeta válida para guardar el consolidado.")
            return
        
        # Deshabilitar botones durante procesamiento
        self.btn_procesar.config(state='disabled')
        self.btn_agregar.config(state='disabled')
        self.btn_carpeta.config(state='disabled')
        self.btn_solo_pdf.config(state='disabled')
        
        # Ejecutar en thread separado para no bloquear UI
        thread = threading.Thread(target=self._procesar_en_background,
                     args=(nombre_salida, formato))
        thread.daemon = True
        thread.start()

    def procesar_solo_pdf(self):
        """Convierte a PDF si es necesario y consolida en un único PDF"""
        if not self.archivos_seleccionados:
            messagebox.showwarning(
                "Advertencia",
                "No hay archivos para procesar.\n\n"
                "Arrastra archivos a la lista o usa los botones para agregar."
            )
            return

        nombre_base = self.entry_nombre.get().strip()
        if not nombre_base:
            nombre_base = "consolidado"

        self._procesar_como_pdf(nombre_base)
        return

        # Determinar carpeta de salida (usar la del primer archivo si no está definida)
        if not self.carpeta_entrada and self.archivos_seleccionados:
            self.carpeta_entrada = os.path.dirname(self.archivos_seleccionados[0])
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, self.carpeta_entrada)
            self.entry_ruta.config(state='readonly')

        # Verificar que hay una ruta de salida válida
        if not self.carpeta_entrada or not os.path.exists(self.carpeta_entrada):
            messagebox.showwarning(
                "Advertencia",
                "Por favor selecciona una carpeta válida para guardar el consolidado."
            )
            return

    def _procesar_como_pdf(self, nombre_base):
        """Valida ruta de salida y lanza el consolidado a PDF"""
        if not DOCX2PDF_AVAILABLE:
            messagebox.showwarning(
                "PDF no disponible",
                "docx2pdf no está instalado. Instálalo para convertir Word a PDF."
            )
            return

        # Determinar carpeta de salida (usar la del primer archivo si no está definida)
        if not self.carpeta_entrada and self.archivos_seleccionados:
            self.carpeta_entrada = os.path.dirname(self.archivos_seleccionados[0])
            self.entry_ruta.config(state='normal')
            self.entry_ruta.delete(0, tk.END)
            self.entry_ruta.insert(0, self.carpeta_entrada)
            self.entry_ruta.config(state='readonly')

        # Verificar que hay una ruta de salida válida
        if not self.carpeta_entrada or not os.path.exists(self.carpeta_entrada):
            messagebox.showwarning(
                "Advertencia",
                "Por favor selecciona una carpeta válida para guardar el consolidado."
            )
            return

        nombre_salida = f"{nombre_base}.pdf"

        # Deshabilitar botones durante procesamiento
        self.btn_procesar.config(state='disabled')
        self.btn_agregar.config(state='disabled')
        self.btn_carpeta.config(state='disabled')
        self.btn_solo_pdf.config(state='disabled')

        # Ejecutar en thread separado para no bloquear UI
        thread = threading.Thread(
            target=self._procesar_solo_pdf_en_background,
            args=(nombre_salida,)
        )
        thread.daemon = True
        thread.start()
    
    def _procesar_en_background(self, nombre_salida, formato_salida):
        """Procesa documentos en background con máxima robustez - Solo consolida contenido"""
        temp_files = []
        try:
            self.progress.start(10)
            self.label_estado.config(text="Estado: Iniciando consolidación...", foreground="blue")
            self.root.update()
            
            # Crear documento Word final
            doc_final = Document()
            
            for idx, ruta_archivo in enumerate(self.archivos_seleccionados, 1):
                nombre_archivo = os.path.basename(ruta_archivo)
                
                try:
                    self.label_estado.config(
                        text=f"Estado: Procesando {idx}/{len(self.archivos_seleccionados)}: {nombre_archivo}",
                        foreground="blue"
                    )
                    self.root.update()
                    
                    if not os.path.exists(ruta_archivo):
                        raise FileNotFoundError(f"El archivo no existe: {ruta_archivo}")
                    
                    # Si está activa la normalización a PDF
                    if self.normalizar_a_pdf.get() and DOCX2PDF_AVAILABLE:
                        # Convertir DOCX a PDF temporal si es necesario
                        if ruta_archivo.lower().endswith(('.docx', '.doc')):
                            # Convertir DOCX a PDF temporal
                            pdf_temporal = os.path.join(self.carpeta_entrada, f"temp_norm_{id(ruta_archivo)}.pdf")
                            temp_files.append(pdf_temporal)
                            try:
                                docx_to_pdf(ruta_archivo, pdf_temporal)
                                ruta_archivo = pdf_temporal  # Usar el PDF temporal
                            except Exception as e:
                                raise Exception(f"Error al convertir DOCX a PDF: {str(e)}")
                        
                        # Procesar como PDF (ya sea original o convertido)
                        self._procesar_pdf(ruta_archivo, doc_final, temp_files, nombre_archivo)
                    
                    else:
                        # Modo normal: procesar según tipo de archivo
                        if ruta_archivo.lower().endswith('.pdf'):
                            self._procesar_pdf(ruta_archivo, doc_final, temp_files, nombre_archivo)
                        
                        elif ruta_archivo.lower().endswith(('.docx', '.doc')):
                            self._procesar_docx(ruta_archivo, doc_final, nombre_archivo)
                    
                    # Agregar separador entre documentos
                    if idx < len(self.archivos_seleccionados):
                        if self.agregar_pagina_blanca.get():
                            # Agregar página en blanco (dos saltos de página)
                            doc_final.add_page_break()
                            doc_final.add_paragraph("")  # Párrafo vacío en página en blanco
                            doc_final.add_page_break()
                        else:
                            # Solo salto de página
                            doc_final.add_page_break()
                        
                except Exception as e:
                    error_msg = f"Error procesando {nombre_archivo}"
                    print(f"ERROR DETALLADO: {error_msg}\n{traceback.format_exc()}")
                    # Agregar notificación de error en el documento
                    self._agregar_error_documento(doc_final, nombre_archivo, str(e))
                    continue
            
            # Guardar documento final
            ruta_salida = os.path.join(self.carpeta_entrada, nombre_salida)
            
            # Intentar guardar múltiples veces en caso de bloqueo de archivo
            max_intentos = 3
            for intento in range(max_intentos):
                try:
                    doc_final.save(ruta_salida)
                    break
                except Exception as e:
                    if intento < max_intentos - 1:
                        import time
                        time.sleep(1)
                    else:
                        raise

            # Convertir a PDF final si se solicitó (misma lógica: PDF se une, Word se convierte a PDF y se une)
            ruta_pdf_final = None
            if formato_salida == "DOCX+PDF":
                try:
                    base, _ = os.path.splitext(ruta_salida)
                    ruta_pdf_final = f"{base}.pdf"

                    pdfs_a_unir = self._preparar_pdfs_para_union(temp_files)
                    if not pdfs_a_unir:
                        raise Exception("No hay PDFs válidos para unir.")

                    self._unir_pdfs(pdfs_a_unir, ruta_pdf_final)
                except Exception as e:
                    messagebox.showwarning(
                        "Error al generar PDF",
                        f"El DOCX se guardó, pero falló la consolidación a PDF:\n{str(e)}"
                    )
            
            self.progress.stop()
            estado_texto = f"Estado: ✅ Consolidación completada: {nombre_salida}"
            if ruta_pdf_final:
                estado_texto += " (PDF generado)"
            
            self.label_estado.config(
                text=estado_texto,
                foreground="green"
            )
            
            detalles = [
                "✅ Documento consolidado guardado exitosamente:",
                f"DOCX: {ruta_salida}",
                f"Total de archivos procesados: {len(self.archivos_seleccionados)}",
                f"Tamaño DOCX: {os.path.getsize(ruta_salida) / 1024:.2f} KB"
            ]
            if ruta_pdf_final and os.path.exists(ruta_pdf_final):
                detalles.append(f"PDF: {ruta_pdf_final}")
            
            messagebox.showinfo("Éxito", "\n".join(detalles))
            
        except Exception as e:
            self.progress.stop()
            error_completo = f"{str(e)}\n\n{traceback.format_exc()}"
            print(f"ERROR CRÍTICO:\n{error_completo}")
            self.label_estado.config(text=f"Estado: ❌ Error crítico", foreground="red")
            messagebox.showerror("Error Crítico", 
                              f"❌ Error al procesar documentos:\n\n{str(e)}\n\n"
                              f"Verifica que los archivos no estén dañados y que Word esté cerrado.")
        
        finally:
            # Limpiar archivos temporales
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass
            
            # Rehabilitar botones
            self.btn_procesar.config(state='normal')
            self.btn_agregar.config(state='normal')
            self.btn_carpeta.config(state='normal')
            self.btn_solo_pdf.config(state='normal')

    def _procesar_solo_pdf_en_background(self, nombre_salida):
        """Consolida PDFs directamente. Si el archivo es Word, lo convierte a PDF y luego une."""
        temp_files = []
        try:
            self.progress.start(10)
            self.label_estado.config(text="Estado: Iniciando consolidación a PDF...", foreground="blue")
            self.root.update()

            pdfs_a_unir = self._preparar_pdfs_para_union(temp_files)
            if not pdfs_a_unir:
                raise Exception("No hay PDFs válidos para unir.")

            ruta_salida = os.path.join(self.carpeta_entrada, nombre_salida)

            # Unir PDFs directamente
            self._unir_pdfs(pdfs_a_unir, ruta_salida)

            self.progress.stop()
            self.label_estado.config(
                text=f"Estado: ✅ Consolidación PDF completada: {nombre_salida}",
                foreground="green"
            )

            detalles = [
                "✅ PDF consolidado guardado exitosamente:",
                f"PDF: {ruta_salida}",
                f"Total de archivos procesados: {len(self.archivos_seleccionados)}",
                f"Tamaño PDF: {os.path.getsize(ruta_salida) / 1024:.2f} KB"
            ]
            messagebox.showinfo("Éxito", "\n".join(detalles))

        except Exception as e:
            self.progress.stop()
            error_completo = f"{str(e)}\n\n{traceback.format_exc()}"
            print(f"ERROR CRÍTICO:\n{error_completo}")
            self.label_estado.config(text=f"Estado: ❌ Error crítico", foreground="red")
            messagebox.showerror(
                "Error Crítico",
                f"❌ Error al procesar PDF:\n\n{str(e)}\n\n"
                f"Verifica que los archivos no estén dañados y que Word esté cerrado."
            )
        finally:
            for temp_file in temp_files:
                try:
                    if os.path.exists(temp_file):
                        os.remove(temp_file)
                except:
                    pass

            self.btn_procesar.config(state='normal')
            self.btn_agregar.config(state='normal')
            self.btn_carpeta.config(state='normal')
            self.btn_solo_pdf.config(state='normal')
    
    def _procesar_pdf(self, ruta_pdf, doc_final, temp_files, nombre_archivo):
        """Procesa un archivo PDF convirtiéndolo a DOCX temporalmente"""
        try:
            # Crear archivo DOCX temporal desde PDF
            temp_docx = os.path.join(self.carpeta_entrada, f"temp_{id(ruta_pdf)}.docx")
            temp_files.append(temp_docx)
            
            try:
                # Convertir PDF a DOCX usando pdf2docx
                converter = Converter(ruta_pdf)
                converter.convert(temp_docx)
                converter.close()
            except Exception as e:
                raise Exception(f"No se puede convertir PDF a DOCX: {str(e)}")
            
            # Cargar el DOCX convertido
            try:
                doc_temp = Document(temp_docx)
            except Exception as e:
                raise Exception(f"No se puede cargar DOCX temporal: {str(e)}")
            
            # Copiar contenido del DOCX al documento final
            try:
                for elemento in doc_temp.element.body:
                    # Copia profunda del elemento para evitar referencias
                    doc_final.element.body.append(deepcopy(elemento))
            except Exception as e:
                print(f"Error en copia directa: {str(e)}")
                # Fallback: copiar tabla por tabla si falla copia directa
                try:
                    for para in doc_temp.paragraphs:
                        if para.text.strip():
                            doc_final.add_paragraph(para.text)
                    
                    for table in doc_temp.tables:
                        new_table = doc_final.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.rows[i].cells[j].text = cell.text
                except Exception as e2:
                    raise Exception(f"Error en fallback: {str(e2)}")
                    
        except Exception as e:
            raise Exception(f"Error al procesar PDF: {str(e)}")
    
    def _procesar_docx(self, ruta_docx, doc_final, nombre_archivo):
        """Procesa un archivo DOCX/DOC - Copia tablas y contenido directamente"""
        try:
            # Intentar cargar documento
            try:
                doc_temp = Document(ruta_docx)
            except Exception as e:
                raise Exception(f"No se puede abrir el documento: {str(e)}")
            
            # Copiar TODOS los elementos del documento (párrafos, tablas, etc.)
            # directamente sin modificaciones
            try:
                for elemento in doc_temp.element.body:
                    # Copia profunda del elemento para evitar referencias
                    doc_final.element.body.append(deepcopy(elemento))
                    
            except Exception as e:
                print(f"Error en copia directa: {str(e)}")
                # Fallback: copiar tabla por tabla
                try:
                    for para in doc_temp.paragraphs:
                        if para.text.strip():
                            doc_final.add_paragraph(para.text)
                    
                    for table in doc_temp.tables:
                        new_table = doc_final.add_table(rows=len(table.rows), cols=len(table.columns))
                        new_table.style = table.style
                        
                        for i, row in enumerate(table.rows):
                            for j, cell in enumerate(row.cells):
                                new_table.rows[i].cells[j].text = cell.text
                except Exception as e2:
                    raise Exception(f"Error en fallback: {str(e2)}")
                    
        except Exception as e:
            raise Exception(f"Error al procesar DOCX: {str(e)}")
    
    def _agregar_error_documento(self, doc_final, nombre_archivo, error_msg):
        """Agrega una notificación de error en el documento"""
        try:
            p_error = doc_final.add_paragraph()
            run_error = p_error.add_run(f"⚠️ ERROR: {nombre_archivo}")
            run_error.bold = True
            run_error.font.color.rgb = RGBColor(255, 0, 0)
            
            p_detalles = doc_final.add_paragraph(f"Razón: {error_msg}")
            p_detalles.style = 'Quote'
        except:
            pass

    def _agregar_error_documento_pdf(self, nombre_archivo, error_msg):
        """Notifica errores durante la preparación de PDFs"""
        try:
            self.label_estado.config(
                text=f"Estado: ⚠️ Error en {nombre_archivo}: {error_msg}",
                foreground="orange"
            )
            self.root.update()
        except:
            pass

    def _preparar_pdfs_para_union(self, temp_files):
        """Convierte Word a PDF cuando sea necesario y devuelve la lista de PDFs a unir."""
        pdfs_a_unir = []

        for idx, ruta_archivo in enumerate(self.archivos_seleccionados, 1):
            nombre_archivo = os.path.basename(ruta_archivo)
            try:
                self.label_estado.config(
                    text=f"Estado: Preparando {idx}/{len(self.archivos_seleccionados)}: {nombre_archivo}",
                    foreground="blue"
                )
                self.root.update()

                if not os.path.exists(ruta_archivo):
                    raise FileNotFoundError(f"El archivo no existe: {ruta_archivo}")

                if ruta_archivo.lower().endswith('.pdf'):
                    pdfs_a_unir.append(ruta_archivo)
                elif ruta_archivo.lower().endswith(('.docx', '.doc')):
                    if not DOCX2PDF_AVAILABLE:
                        raise Exception("docx2pdf no está instalado. Instálalo para convertir Word a PDF.")

                    pdf_temporal = os.path.join(
                        self.carpeta_entrada,
                        f"temp_conv_{id(ruta_archivo)}.pdf"
                    )
                    temp_files.append(pdf_temporal)
                    try:
                        docx_to_pdf(ruta_archivo, pdf_temporal)
                    except Exception as e:
                        raise Exception(f"Error al convertir Word a PDF: {str(e)}")

                    pdfs_a_unir.append(pdf_temporal)
                else:
                    raise Exception("Formato no soportado para PDF. Usa PDF o Word.")

            except Exception as e:
                error_msg = f"Error preparando {nombre_archivo}"
                print(f"ERROR DETALLADO: {error_msg}\n{traceback.format_exc()}")
                self._agregar_error_documento_pdf(nombre_archivo, str(e))
                continue

        return pdfs_a_unir

    def _unir_pdfs(self, pdfs_a_unir, ruta_salida):
        """Une una lista de PDFs en un solo archivo."""
        pdf_salida = fitz.open()
        try:
            for idx, pdf_path in enumerate(pdfs_a_unir, 1):
                self.label_estado.config(
                    text=f"Estado: Uniendo {idx}/{len(pdfs_a_unir)}: {os.path.basename(pdf_path)}",
                    foreground="blue"
                )
                self.root.update()

                last_rect = None
                with fitz.open(pdf_path) as src:
                    if src.page_count > 0:
                        last_rect = src[-1].rect
                    pdf_salida.insert_pdf(src)

                if idx < len(pdfs_a_unir) and self.agregar_pagina_blanca.get():
                    if last_rect:
                        pdf_salida.new_page(width=last_rect.width, height=last_rect.height)
                    else:
                        pdf_salida.new_page()

            pdf_salida.save(ruta_salida)
        finally:
            pdf_salida.close()


def main():
    if DRAG_DROP_AVAILABLE:
        root = TkinterDnD.Tk()
    else:
        root = tk.Tk()
        messagebox.showinfo(
            "Información",
            "⚠️ La funcionalidad de arrastrar y soltar no está disponible.\n\n"
            "Para habilitarla, instala: pip install tkinterdnd2\n\n"
            "Puedes usar los botones para agregar archivos."
        )
    
    app = ConsolidadorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()

"""
