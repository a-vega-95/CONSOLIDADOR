import os
import sys
import tempfile
import shutil

ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
if ROOT_DIR not in sys.path:
    sys.path.insert(0, ROOT_DIR)

import fitz
from docx import Document

from core import ConsolidationService, DOCX2PDF_AVAILABLE


def _crear_docx(ruta):
    doc = Document()
    doc.add_heading("Prueba DOCX", level=1)
    doc.add_paragraph("Contenido de prueba para consolidación.")
    doc.save(ruta)


def _crear_pdf(ruta):
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 72), "PDF de prueba")
    pdf.save(ruta)
    pdf.close()


def main():
    temp_dir = tempfile.mkdtemp(prefix="consolidador_test_")
    try:
        docx_path = os.path.join(temp_dir, "1.docx")
        pdf_path = os.path.join(temp_dir, "2.pdf")

        _crear_docx(docx_path)
        _crear_pdf(pdf_path)

        service = ConsolidationService()

        # Prueba de consolidación a DOCX
        salida_docx = os.path.join(temp_dir, "salida.docx")
        resultado_docx = service.consolidate_to_docx(
            [docx_path, pdf_path],
            salida_docx,
            add_blank_page=False,
            normalize_to_pdf=False,
            generate_pdf=False,
        )
        print("DOCX generado:", resultado_docx["docx"], os.path.exists(resultado_docx["docx"]))

        # Prueba de consolidación a PDF (si docx2pdf está disponible)
        if DOCX2PDF_AVAILABLE:
            salida_pdf = os.path.join(temp_dir, "salida.pdf")
            resultado_pdf = service.consolidate_to_pdf(
                [docx_path, pdf_path],
                salida_pdf,
                add_blank_page=False,
            )
            print("PDF generado:", resultado_pdf, os.path.exists(resultado_pdf))
        else:
            print("docx2pdf no disponible: se omite prueba de PDF")

        print("✅ Prueba rápida completada")
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


if __name__ == "__main__":
    main()
