import os
import traceback
from copy import deepcopy

import fitz  # PyMuPDF
from docx import Document
from docx.shared import RGBColor
from pdf2docx.converter import Converter

try:
    from docx2pdf import convert as docx_to_pdf
    DOCX2PDF_AVAILABLE = True
except ImportError:
    docx_to_pdf = None
    DOCX2PDF_AVAILABLE = False


class ConsolidationService:
    def __init__(self):
        self.docx2pdf_available = DOCX2PDF_AVAILABLE
        self.docx_to_pdf = docx_to_pdf

    def consolidate_to_docx(
        self,
        files,
        output_docx,
        add_blank_page=False,
        normalize_to_pdf=False,
        generate_pdf=False,
        on_status=None,
        on_item_error=None,
    ):
        temp_files = []
        doc_final = Document()

        for idx, ruta_archivo in enumerate(files, 1):
            nombre_archivo = os.path.basename(ruta_archivo)
            try:
                self._call_status(
                    on_status,
                    f"Estado: Procesando {idx}/{len(files)}: {nombre_archivo}",
                    "blue",
                )

                if not os.path.exists(ruta_archivo):
                    raise FileNotFoundError(f"El archivo no existe: {ruta_archivo}")

                if normalize_to_pdf and self.docx2pdf_available:
                    if ruta_archivo.lower().endswith((".docx", ".doc")):
                        pdf_temporal = os.path.join(
                            os.path.dirname(output_docx),
                            f"temp_norm_{id(ruta_archivo)}.pdf",
                        )
                        temp_files.append(pdf_temporal)
                        try:
                            self.docx_to_pdf(ruta_archivo, pdf_temporal)
                            ruta_archivo = pdf_temporal
                        except Exception as e:
                            raise Exception(f"Error al convertir DOCX a PDF: {str(e)}")

                    self._process_pdf(ruta_archivo, doc_final, temp_files)
                else:
                    if ruta_archivo.lower().endswith(".pdf"):
                        self._process_pdf(ruta_archivo, doc_final, temp_files)
                    elif ruta_archivo.lower().endswith((".docx", ".doc")):
                        self._process_docx(ruta_archivo, doc_final)
                    else:
                        raise Exception("Formato no soportado. Usa PDF o Word.")

                if idx < len(files):
                    if add_blank_page:
                        doc_final.add_page_break()
                        doc_final.add_paragraph("")
                        doc_final.add_page_break()
                    else:
                        doc_final.add_page_break()

            except Exception as e:
                error_msg = f"Error procesando {nombre_archivo}"
                print(f"ERROR DETALLADO: {error_msg}\n{traceback.format_exc()}")
                self._add_error_to_doc(doc_final, nombre_archivo, str(e))
                self._call_item_error(on_item_error, nombre_archivo, str(e))
                continue

        self._safe_save_docx(doc_final, output_docx)

        pdf_path = None
        if generate_pdf:
            if not self.docx2pdf_available:
                raise Exception("docx2pdf no está instalado. No se puede generar PDF.")

            base, _ = os.path.splitext(output_docx)
            pdf_path = f"{base}.pdf"
            pdfs_a_unir = self.prepare_pdfs_for_merge(
                files,
                os.path.dirname(output_docx),
                temp_files,
                on_status=on_status,
                on_item_error=on_item_error,
            )
            if not pdfs_a_unir:
                raise Exception("No hay PDFs válidos para unir.")
            self.merge_pdfs(pdfs_a_unir, pdf_path, add_blank_page, on_status=on_status)

        self._cleanup_temp_files(temp_files)
        return {"docx": output_docx, "pdf": pdf_path}

    def consolidate_to_pdf(
        self,
        files,
        output_pdf,
        add_blank_page=False,
        on_status=None,
        on_item_error=None,
    ):
        temp_files = []
        pdfs_a_unir = self.prepare_pdfs_for_merge(
            files,
            os.path.dirname(output_pdf),
            temp_files,
            on_status=on_status,
            on_item_error=on_item_error,
        )
        if not pdfs_a_unir:
            raise Exception("No hay PDFs válidos para unir.")

        self.merge_pdfs(pdfs_a_unir, output_pdf, add_blank_page, on_status=on_status)
        self._cleanup_temp_files(temp_files)
        return output_pdf

    def prepare_pdfs_for_merge(
        self,
        files,
        output_dir,
        temp_files,
        on_status=None,
        on_item_error=None,
    ):
        pdfs_a_unir = []

        for idx, ruta_archivo in enumerate(files, 1):
            nombre_archivo = os.path.basename(ruta_archivo)
            try:
                self._call_status(
                    on_status,
                    f"Estado: Preparando {idx}/{len(files)}: {nombre_archivo}",
                    "blue",
                )

                if not os.path.exists(ruta_archivo):
                    raise FileNotFoundError(f"El archivo no existe: {ruta_archivo}")

                if ruta_archivo.lower().endswith(".pdf"):
                    pdfs_a_unir.append(ruta_archivo)
                elif ruta_archivo.lower().endswith((".docx", ".doc")):
                    if not self.docx2pdf_available:
                        raise Exception("docx2pdf no está instalado. Instálalo para convertir Word a PDF.")

                    pdf_temporal = os.path.join(
                        output_dir,
                        f"temp_conv_{id(ruta_archivo)}.pdf",
                    )
                    temp_files.append(pdf_temporal)
                    try:
                        self.docx_to_pdf(ruta_archivo, pdf_temporal)
                    except Exception as e:
                        raise Exception(f"Error al convertir Word a PDF: {str(e)}")

                    pdfs_a_unir.append(pdf_temporal)
                else:
                    raise Exception("Formato no soportado para PDF. Usa PDF o Word.")

            except Exception as e:
                error_msg = f"Error preparando {nombre_archivo}"
                print(f"ERROR DETALLADO: {error_msg}\n{traceback.format_exc()}")
                self._call_item_error(on_item_error, nombre_archivo, str(e))
                continue

        return pdfs_a_unir

    def merge_pdfs(self, pdfs, output_pdf, add_blank_page=False, on_status=None):
        pdf_salida = fitz.open()
        try:
            for idx, pdf_path in enumerate(pdfs, 1):
                self._call_status(
                    on_status,
                    f"Estado: Uniendo {idx}/{len(pdfs)}: {os.path.basename(pdf_path)}",
                    "blue",
                )

                last_rect = None
                with fitz.open(pdf_path) as src:
                    if src.page_count > 0:
                        last_rect = src[-1].rect
                    pdf_salida.insert_pdf(src)

                if idx < len(pdfs) and add_blank_page:
                    if last_rect:
                        pdf_salida.new_page(width=last_rect.width, height=last_rect.height)
                    else:
                        pdf_salida.new_page()

            pdf_salida.save(output_pdf)
        finally:
            pdf_salida.close()

    def _process_pdf(self, ruta_pdf, doc_final, temp_files):
        temp_docx = os.path.join(os.path.dirname(ruta_pdf), f"temp_{id(ruta_pdf)}.docx")
        temp_files.append(temp_docx)

        try:
            converter = Converter(ruta_pdf)
            converter.convert(temp_docx)
            converter.close()
        except Exception as e:
            raise Exception(f"No se puede convertir PDF a DOCX: {str(e)}")

        try:
            doc_temp = Document(temp_docx)
        except Exception as e:
            raise Exception(f"No se puede cargar DOCX temporal: {str(e)}")

        try:
            for elemento in doc_temp.element.body:
                doc_final.element.body.append(deepcopy(elemento))
        except Exception as e:
            print(f"Error en copia directa: {str(e)}")
            try:
                for para in doc_temp.paragraphs:
                    if para.text.strip():
                        doc_final.add_paragraph(para.text)

                for table in doc_temp.tables:
                    new_table = doc_final.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
            except Exception as e2:
                raise Exception(f"Error en fallback: {str(e2)}")

    def _process_docx(self, ruta_docx, doc_final):
        try:
            doc_temp = Document(ruta_docx)
        except Exception as e:
            raise Exception(f"No se puede abrir el documento: {str(e)}")

        try:
            for elemento in doc_temp.element.body:
                doc_final.element.body.append(deepcopy(elemento))
        except Exception as e:
            print(f"Error en copia directa: {str(e)}")
            try:
                for para in doc_temp.paragraphs:
                    if para.text.strip():
                        doc_final.add_paragraph(para.text)

                for table in doc_temp.tables:
                    new_table = doc_final.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            new_table.rows[i].cells[j].text = cell.text
            except Exception as e2:
                raise Exception(f"Error en fallback: {str(e2)}")

    def _add_error_to_doc(self, doc_final, nombre_archivo, error_msg):
        try:
            p_error = doc_final.add_paragraph()
            run_error = p_error.add_run(f"⚠️ ERROR: {nombre_archivo}")
            run_error.bold = True
            run_error.font.color.rgb = RGBColor(255, 0, 0)

            p_detalles = doc_final.add_paragraph(f"Razón: {error_msg}")
            p_detalles.style = "Quote"
        except Exception:
            pass

    def _safe_save_docx(self, doc_final, output_docx):
        max_intentos = 3
        for intento in range(max_intentos):
            try:
                doc_final.save(output_docx)
                break
            except Exception:
                if intento < max_intentos - 1:
                    import time

                    time.sleep(1)
                else:
                    raise

    def _cleanup_temp_files(self, temp_files):
        for temp_file in temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except Exception:
                pass

    def _call_status(self, on_status, message, color):
        if on_status:
            on_status(message, color)

    def _call_item_error(self, on_item_error, nombre_archivo, error_msg):
        if on_item_error:
            on_item_error(nombre_archivo, error_msg)
